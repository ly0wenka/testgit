from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
from pathlib import Path

folder = Path(r"S:\Users\L\Downloads\New folder (175)\testgit\os\l8\imgs")
doc = Document()

# Заголовок
heading = doc.add_heading("Зображення з теки imgs", level=1)
run = heading.runs[0]
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0,0,0)

# Стиль підписів
caption_style = "Caption"
if caption_style not in doc.styles:
    if "Підпис" in doc.styles:
        caption_style = "Підпис"
    else:
        caption_style = None

exts = {".png", ".jpg", ".jpeg", ".gif", ".bmp"}
index = 1

for img in sorted(folder.iterdir()):
    if img.suffix.lower() in exts:

        # ---------------------------
        # 1. Текст-посилання
        # ---------------------------
        reference_text = f"На рис. {index} зображено {img.stem}."
        ref_p = doc.add_paragraph(reference_text)
        ref_p.alignment = 0  # ліворуч
        ref_p.paragraph_format.keep_with_next = True  # тримати з наступним абзацом
        run = ref_p.runs[0]
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0,0,0)

        # ---------------------------
        # 2. Вставка зображення
        # ---------------------------
        pic = doc.add_picture(str(img), width=Inches(6))
        last_paragraph = doc.paragraphs[-1]  # абзац із картинкою
        last_paragraph.paragraph_format.keep_with_next = True  # тримати з підписом

        # ---------------------------
        # 3. Підпис під малюнком
        # ---------------------------
        caption_text = f"Рис. {index} — {img.stem}"
        caption_paragraph = doc.add_paragraph(caption_text)
        if caption_style:
            caption_paragraph.style = caption_style
        caption_paragraph.alignment = 1  # центр
        caption_paragraph.paragraph_format.keep_with_next = False  # останній абзац блоку
        run = caption_paragraph.runs[0]
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0,0,0)

        # невеликий відступ
        doc.add_paragraph()

        index += 1

# Збереження
output_path = folder / "imgs.docx"
doc.save(output_path)
print("Готово:", output_path)
