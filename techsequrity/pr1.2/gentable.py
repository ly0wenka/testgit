from docx import Document

# Create document
doc = Document()
doc.add_heading('Звіт про кіберінциденти (SC Media)', level=1)

incidents = [
    {
        "title": "1. European Commission – Ivanti EPMM Vulnerabilities (2026)",
        "fields": {
            "Дата інциденту": "30 січня 2026 р.",
            "Постраждала організація": "Європейська Комісія (European Commission)",
            "Кількість постраждалих": "Не вказано; потенційний доступ до контактних даних співробітників",
            "Що було зроблено?": "CERT-EU локалізував інцидент (~9 годин), очищено системи, випущені патчі Ivanti.",
            "Які експлойти використовувалися?": "CVE-2026-1281 та CVE-2026-1340 (RCE, code injection, Ivanti EPMM).",
            "Яким чином налагоджено захист?": "Негайний патчинг, сегментація доступу, MFA, моніторинг логів, аудит облікових записів.",
            "Посилання": "https://www.scworld.com/brief/european-commission-hit-by-cyberattack-linked-to-ivanti-software-flaws"
        }
    },
    {
        "title": "2. France Travail – Data Breach (2024, fine in 2026)",
        "fields": {
            "Дата інциденту": "Початок 2024 р.; штраф 22 січня 2026 р.",
            "Постраждала організація": "France Travail (Національне агентство зайнятості Франції)",
            "Кількість постраждалих": "43 мільйони осіб",
            "Що було зроблено?": "CNIL наклав штраф €5 млн; зобов'язання посилити захист (щоденний штраф за невиконання).",
            "Які експлойти використовувалися?": "Соціальна інженерія, компрометація облікових записів партнерської організації.",
            "Яким чином налагоджено захист?": "MFA, принцип мінімальних привілеїв, моніторинг логів, навчання персоналу.",
            "Посилання": "https://www.scworld.com/brief/france-fines-employment-agency-e5-million-for-data-breach-affecting-43-million"
        }
    },
    {
        "title": "3. Scattered Spider – Rapid Ransomware Deployment (2025)",
        "fields": {
            "Дата інциденту": "2025 р. (звіт CrowdStrike)",
            "Постраждала організація": "Не конкретизовано (загальний тренд атак)",
            "Кількість постраждалих": "Не вказано",
            "Що було зроблено?": "Атаки скорочені до 24 годин від доступу до шифрування; використання SaaS і Entra ID.",
            "Які експлойти використовувалися?": "Vishing, компрометація MFA, living-off-the-land техніки.",
            "Яким чином налагоджено захист?": "Посилення IAM, MFA, поведінковий моніторинг, сегментація доступу.",
            "Посилання": "https://www.scworld.com/news/scattered-spider-now-moves-from-initial-access-to-encryption-in-24-hours"
        }
    }
]

for incident in incidents:
    doc.add_heading(incident["title"], level=2)
    table = doc.add_table(rows=len(incident["fields"]), cols=2)
    table.style = 'Table Grid'
    
    for i, (key, value) in enumerate(incident["fields"].items()):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value

    doc.add_paragraph("")

doc.save("Звіт_про_кіберінциденти_SC_Media.docx")

print("Файл створено успішно!")
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Створення документа
doc = Document()
doc.add_heading('Звіт про шкідливе ПЗ проти українських цілей (CoreWin)', level=1)

# Дані інциденту
incident = {
    "Дата інциденту": "Січень–квітень 2022 року",
    "Постраждала організація": "Державні органи України, банки, енергетика, телеком",
    "Кількість постраждалих": "Десятки державних сайтів, сотні систем, тисячі пристроїв (KA-SAT)",
    "Що було зроблено?": (
        "Здійснено масові кібератаки із застосуванням destructive malware "
        "(вайпери) для знищення даних, дефейсу сайтів та атак на критичну інфраструктуру."
    ),
    "Які експлойти / шкідливе ПЗ використовувалися?": (
        "WhisperGate, HermeticWiper, IsaacWiper, AcidRain, CaddyWiper, "
        "DoubleZero, Industroyer2."
    ),
    "Яким чином налагоджено захист?": (
        "Сегментація мереж (IT/OT), резервне копіювання, EDR/AV рішення, "
        "MFA, моніторинг логів, навчання персоналу, аудит доступів."
    ),
    "Посилання на джерело": "https://corewin.ua/blog/malware-being-used-on-ukrainian-targets/"
}

# Додавання таблиці
table = doc.add_table(rows=len(incident), cols=2)
table.style = 'Table Grid'

for i, (key, value) in enumerate(incident.items()):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = value

# Збереження файлу
doc.save("Звіт_CoreWin_Malware_UA.docx")

print("DOCX файл успішно створено!")
from docx import Document

# Створення документа
doc = Document()
doc.add_heading('Звіт про кіберінцидент (SC Media)', level=1)

incident = {
    "Дата інциденту": "З середини 2024 р.; оприлюднено у 2026 р.",
    "Постраждала організація": "Різні організації, що використовують Dell RecoverPoint for VMs",
    "Кількість постраждалих": "Не розкрито (цільові атаки на корпоративні середовища)",
    "Що було зроблено?": (
        "Кіберугруповання UNC6201, пов’язане з Китаєм, експлуатувало "
        "уразливість у Dell RecoverPoint for VMs версії 10.0 для отримання "
        "стійкого доступу, збору даних та подальшого переміщення в мережі."
    ),
    "Які експлойти використовувалися?": (
        "Експлуатація невідомої (ймовірно 0-day) уразливості в Dell RecoverPoint "
        "for VMs 10.0 для віддаленого доступу та закріплення в системі."
    ),
    "Яким чином налагоджено захист?": (
        "Оновлення та патчинг Dell RecoverPoint, ізоляція керуючих інтерфейсів, "
        "використання MFA, моніторинг логів, контроль доступу та сегментація мережі."
    ),
    "Посилання": "https://www.scworld.com/news/china-linked-unc6201-exploits-10-0-bug-in-dell-recoverpoint-for-vms-since-mid-2024"
}

# Додавання таблиці
table = doc.add_table(rows=len(incident), cols=2)
table.style = 'Table Grid'

for i, (key, value) in enumerate(incident.items()):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = value

# Збереження файлу
doc.save("Звіт_UNC6201_Dell_RecoverPoint.docx")

print("DOCX файл успішно створено!")
from docx import Document

# Створення документа
doc = Document()
doc.add_heading('Звіт про кіберінцидент (SC Media)', level=1)

incident = {
    "Дата інциденту": "Лютий 2026 р. (опубліковано дослідження)",
    "Постраждала організація": "Користувачі Visual Studio Code та сумісних IDE (Cursor, Windsurf тощо)",
    "Кількість постраждалих": (
        "Потенційно мільйони розробників; вразливі розширення мають понад 125–128 млн завантажень"
    ),
    "Що було зроблено?": (
        "Дослідники виявили критичні вразливості в популярних розширеннях VS Code, "
        "які дозволяють виконання довільного коду та викрадення локальних файлів. "
        "Атаки можливі через відкриття шкідливих конфігурацій, HTML-файлів або робочих просторів."
    ),
    "Які експлойти використовувалися?": (
        "CVE-2025-65717 (Live Server), CVE-2025-65715 (Code Runner), "
        "CVE-2025-65716 (Markdown Preview Enhanced); XSS та небезпечні налаштування workspace."
    ),
    "Яким чином налагоджено захист?": (
        "Оновлення та видалення вразливих розширень, використання лише перевірених плагінів, "
        "обмеження відкриття недовірених файлів/конфігурацій, моніторинг змін налаштувань IDE, "
        "застосування принципу мінімальних привілеїв і контроль доступу до локальних ресурсів."
    ),
    "Посилання": "https://www.scworld.com/brief/critical-vscode-extension-vulnerabilities-could-lead-to-code-execution-and-data-theft"
}

# Додавання таблиці
table = doc.add_table(rows=len(incident), cols=2)
table.style = 'Table Grid'

for i, (key, value) in enumerate(incident.items()):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = value

# Збереження
doc.save("Звіт_VSCode_Extension_Vulnerabilities.docx")

print("DOCX файл успішно створено!")
from docx import Document

# Створення документа
doc = Document()
doc.add_heading('Звіт про кіберінцидент (SC Media)', level=1)

incident = {
    "Дата інциденту": "Лютий 2026 р.",
    "Постраждала організація": "Користувачі Google Chrome (Windows, macOS, Linux)",
    "Кількість постраждалих": (
        "Потенційно мільярди користувачів Chrome; активна експлуатація zero-day у дикій природі"
    ),
    "Що було зроблено?": (
        "Google випустила екстрене оновлення Chrome для усунення активно "
        "експлуатованої zero-day уразливості, що дозволяє виконання довільного коду "
        "через спеціально сформовані вебсторінки."
    ),
    "Які експлойти використовувалися?": (
        "CVE-2026-2441 — use-after-free у CSS компоненті Chrome; "
        "можливе виконання довільного коду в sandbox через шкідливу HTML-сторінку."
    ),
    "Яким чином налагоджено захист?": (
        "Негайне оновлення браузера до Chrome 145.0.7632.75/76 (Win/Mac) "
        "або 144.0.7559.75 (Linux), застосування автоматичних оновлень, "
        "моніторинг підозрілих вебресурсів, використання sandbox-захисту та EDR."
    ),
    "Посилання": "https://www.scworld.com/brief/google-releases-emergency-chrome-update-for-zero-day-exploit"
}

# Додавання таблиці
table = doc.add_table(rows=len(incident), cols=2)
table.style = 'Table Grid'

for i, (key, value) in enumerate(incident.items()):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = value

# Збереження
doc.save("Звіт_Chrome_ZeroDay_Exploit.docx")

print("DOCX файл успішно створено!")
from docx import Document

# Створення документа
doc = Document()
doc.add_heading('Звіт про кіберінцидент (SC Media)', level=1)

incident = {
    "Дата інциденту": "Лютий 2026 р.",
    "Постраждала організація": (
        "WordPress-сайти, що використовують плагін WPvivid Backup & Migration"
    ),
    "Кількість постраждалих": (
        "Понад 900 000 сайтів WordPress, що використовують вразливі версії плагіна"
    ),
    "Що було зроблено?": (
        "Виявлено критичну вразливість, яка дозволяє неавторизованим зловмисникам "
        "завантажувати довільні файли та виконувати віддалений код на сервері, "
        "що може призвести до повного захоплення сайту."
    ),
    "Які експлойти використовувалися?": (
        "CVE-2026-1357 — Unauthenticated Arbitrary File Upload, що призводить до RCE "
        "через помилки обробки RSA-дешифрування та відсутність санітизації шляхів."
    ),
    "Яким чином налагоджено захист?": (
        "Оновлення плагіна до версії 0.9.124 або новішої, вимкнення функції "
        "'receive backup from another site', застосування WAF, контроль завантажень, "
        "моніторинг логів та перевірка сайту на наявність вебшелів."
    ),
    "Посилання": "https://www.scworld.com/brief/critical-vulnerability-in-wpvivid-backup-plugin-allows-remote-code-execution"
}

# Додавання таблиці
table = doc.add_table(rows=len(incident), cols=2)
table.style = 'Table Grid'

for i, (key, value) in enumerate(incident.items()):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = value

# Збереження файлу
doc.save("Звіт_WPvivid_RCE_Vulnerability.docx")

print("DOCX файл успішно створено!")
from docx import Document

# Створення документа
doc = Document()
doc.add_heading('Звіт про кіберінцидент (SC Media)', level=1)

incident = {
    "Дата інциденту": "Лютий 2026 р. (кампанія тривала майже десятиліття)",
    "Постраждала організація": (
        "Державні установи та телекомунікаційні компанії у багатьох країнах світу"
    ),
    "Кількість постраждалих": (
        "Щонайменше 53 організації у 42 країнах; активність зафіксована майже у 70 країнах"
    ),
    "Що було зроблено?": (
        "Google Threat Intelligence Group (GTIG) разом з Mandiant та партнерами "
        "виявили та зупинили масштабну кампанію кіберрозвідки, яку проводила "
        "група UNC2814, пов’язана з Китаєм. Було відключено інфраструктуру "
        "зловмисників, заблоковано їхні облікові записи та проєкти в Google Cloud."
    ),
    "Які експлойти / шкідливе ПЗ використовувалися?": (
        "Бекдор GRIDTIDE, який використовував Google Sheets API як канал "
        "керування (C2). Команди передавались через таблиці Google Sheets, "
        "що дозволяло маскувати шкідливий трафік під звичайний хмарний трафік."
    ),
    "Яким чином налагоджено захист?": (
        "Виявлення та відключення інфраструктури атакуючих, відкликання доступів "
        "до Google API, моніторинг SaaS-активності, контроль доступу, "
        "сегментація мереж та використання індикаторів компрометації (IoC)."
    ),
    "Посилання": "https://www.scworld.com/news/google-disrupts-decade-long-china-linked-unc2814-espionage-campaign"
}

# Таблиця
table = doc.add_table(rows=len(incident), cols=2)
table.style = 'Table Grid'

for i, (key, value) in enumerate(incident.items()):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = value

doc.save("Звіт_UNC2814_Espionage_Campaign.docx")

print("DOCX файл успішно створено!")

from docx import Document

# Створення документа
doc = Document()
doc.add_heading('Звіт про кіберінцидент (SC Media)', level=1)

incident = {
    "Дата інциденту": "Березень 2026 р.",
    "Постраждала організація": (
        "Організації, що використовують Windows 11 Enterprise з RRAS (Routing and Remote Access Service)"
    ),
    "Кількість постраждалих": (
        "Не розкрито; потенційно корпоративні середовища з використанням RRAS і Windows Autopatch"
    ),
    "Що було зроблено?": (
        "Microsoft випустила позапланове (out-of-band) оновлення KB5084597 для усунення "
        "критичних вразливостей у службі RRAS. Оновлення розгортається як hotpatch "
        "без необхідності перезавантаження систем."
    ),
    "Які експлойти використовувалися?": (
        "CVE-2026-25172, CVE-2026-25173, CVE-2026-26111 — уразливості RRAS, що дозволяють "
        "віддалене виконання коду (RCE) при підключенні до шкідливого сервера."
    ),
    "Яким чином налагоджено захист?": (
        "Встановлення оновлення KB5084597, використання Windows Autopatch, "
        "обмеження підключень до недовірених серверів, моніторинг мережевої активності, "
        "сегментація доступу та контроль прав користувачів."
    ),
    "Посилання": "https://www.scworld.com/brief/microsoft-releases-out-of-band-update-for-windows-11-rras-vulnerabilities"
}

# Додавання таблиці
table = doc.add_table(rows=len(incident), cols=2)
table.style = 'Table Grid'

for i, (key, value) in enumerate(incident.items()):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = value

# Збереження
doc.save("Звіт_Windows11_RRAS_Vulnerabilities.docx")

print("DOCX файл успішно створено!")

from docx import Document

# Створення документа
doc = Document()
doc.add_heading('Звіт про кіберінцидент / тенденцію (SC Media)', level=1)

incident = {
    "Дата інциденту": "Березень 2026 р. (конференція BSides SF)",
    "Постраждала організація": (
        "Глобальна економіка та організації, що впроваджують AI (загальна тенденція)"
    ),
    "Кількість постраждалих": (
        "Не конкретизовано; потенційно мільйони працівників через автоматизацію та скорочення робочих місць"
    ),
    "Що було зроблено?": (
        "На конференції BSides SF експерт з безпеки Кеті Муссуріс попередила, "
        "що неконтрольоване впровадження штучного інтелекту може призвести до "
        "масової концентрації багатства та економічної нерівності."
    ),
    "Які ризики / загрози розглядаються?": (
        "Концентрація влади у великих технологічних компаніях, масова автоматизація, "
        "звільнення працівників, зростання витрат (наприклад, енергоспоживання), "
        "залежність від AI та відсутність регуляторного контролю."
    ),
    "Яким чином налагоджено захист?": (
        "Необхідність державного регулювання AI, участь експертів у формуванні політик, "
        "впровадження стандартів безпеки, контроль впливу AI на ринок праці та економіку, "
        "залучення суспільства до прийняття рішень щодо розвитку технологій."
    ),
    "Додаткова інформація": (
        "AI може стати потужним інструментом для розвитку суспільства, але без контролю "
        "ризикує посилити економічну нерівність та призвести до швидкої концентрації багатства."
    ),
    "Посилання": "https://www.scworld.com/news/bsides-sf-ai-must-benefit-everyone-not-just-the-wealthy"
}

# Додавання таблиці
table = doc.add_table(rows=len(incident), cols=2)
table.style = 'Table Grid'

for i, (key, value) in enumerate(incident.items()):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = value

# Збереження файлу
doc.save("Звіт_BSides_SF_AI_Risk.docx")

print("DOCX файл успішно створено!")

from docx import Document

# Створення документа
doc = Document()
doc.add_heading('Звіт про кіберінцидент (SC Media)', level=1)

incident = {
    "Дата інциденту": "23–24 березня 2026 р.",
    "Постраждала організація": (
        "Компанії, що використовують Citrix NetScaler ADC та NetScaler Gateway"
    ),
    "Кількість постраждалих": (
        "Не вказано; потенційно тисячі корпоративних середовищ по всьому світу"
    ),
    "Що було зроблено?": (
        "Citrix випустила термінові оновлення безпеки для критичної уразливості "
        "в NetScaler ADC після попереджень дослідників про швидке появлення експлойтів."
    ),
    "Які експлойти використовувалися?": (
        "CVE-2026-3055 — out-of-bounds read (витік пам’яті, CVSS 9.3); "
        "CVE-2026-4368 — race condition (змішування сесій користувачів)."
    ),
    "Суть уразливості": (
        "Неавторизований віддалений доступ може дозволити зловмисникам "
        "читати чутливі дані з пам’яті системи, включаючи токени сесій та облікові дані."
    ),
    "Яким чином налагоджено захист?": (
        "Негайне встановлення патчів, оновлення NetScaler до безпечних версій, "
        "перевірка конфігурації SAML IDP, моніторинг логів, сегментація доступу."
    ),
    "Додаткові ризики": (
        "Експерти попереджають, що експлуатація майже гарантовано почнеться "
        "після публікації PoC-коду; уразливість подібна до CitrixBleed."
    ),
    "Посилання": "https://www.scworld.com/news/citrix-patches-critical-netscaler-adc-bug"
}

# Додавання таблиці
table = doc.add_table(rows=len(incident), cols=2)
table.style = 'Table Grid'

for i, (key, value) in enumerate(incident.items()):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = value

# Збереження файлу
doc.save("Звіт_Citrix_NetScaler_Critical_Bug.docx")

print("DOCX файл успішно створено!")

from docx import Document

# Створення документа
doc = Document()
doc.add_heading('Звіт про кіберінцидент (SC Media)', level=1)

incident = {
    "Дата інциденту": "Березень 2026 р. (FBI alert)",
    "Постраждала організація": (
        "Журналісти, дисиденти, опозиційні групи та організації по всьому світу"
    ),
    "Кількість постраждалих": (
        "Не вказано; кампанія активна щонайменше з 2023 року"
    ),
    "Що було зроблено?": (
        "FBI оприлюднило попередження про діяльність іранських кіберугруповань "
        "(зокрема Handala), які використовують Telegram як інфраструктуру для атак."
    ),
    "Які експлойти / техніки використовувалися?": (
        "Соціальна інженерія, фішинг, маскування malware під легітимні програми "
        "(Telegram, KeePass, AI-сервіси), багатоступенева інфекція."
    ),
    "Суть атаки": (
        "Після встановлення шкідливого ПЗ заражений пристрій підключається до "
        "Telegram-бота (C2), що дозволяє зловмисникам отримати віддалений доступ, "
        "викрадати файли, робити скріншоти та стежити за активністю користувача."
    ),
    "Яким чином налагоджено захист?": (
        "Навчання користувачів, використання EDR/AV, контроль доступу до месенджерів, "
        "моніторинг мережевого трафіку, блокування підозрілих файлів і посилань."
    ),
    "Додаткові ризики": (
        "Використання легітимної платформи Telegram як C2 дозволяє маскувати трафік "
        "та обходити системи виявлення загроз."
    ),
    "Посилання": "https://www.scworld.com/news/iran-backed-handala-uses-telegram-for-c2-to-push-malware-fbi-says"
}

# Таблиця
table = doc.add_table(rows=len(incident), cols=2)
table.style = 'Table Grid'

for i, (key, value) in enumerate(incident.items()):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = value

# Збереження
doc.save("Звіт_Handala_Telegram_C2.docx")

print("DOCX файл успішно створено!")

from docx import Document

# Створення документа
doc = Document()
doc.add_heading('Звіт про кіберінцидент (SC Media)', level=1)

incident = {
    "Дата інциденту": "Березень 2026 р.",
    "Постраждала організація": (
        "Користувачі macOS (особливо розробники, користувачі AI-інструментів)"
    ),
    "Кількість постраждалих": (
        "Не вказано; кампанії глобальні, поширюються через рекламу та пошукові системи"
    ),
    "Що було зроблено?": (
        "Дослідники виявили ClickFix кампанії, які поширюють macOS-шкідливе ПЗ "
        "MacSync через фальшиві сторінки, рекламу та інструкції ввести команди в Terminal."
    ),
    "Які техніки використовувалися?": (
        "ClickFix (social engineering), malvertising, SEO poisoning, "
        "фальшиві сайти та GitHub-репозиторії, copy-paste команд у Terminal."
    ),
    "Суть атаки": (
        "Користувачеві пропонують 'виправити проблему' або встановити інструмент, "
        "після чого він сам запускає шкідливу команду в Terminal, яка завантажує MacSync."
    ),
    "Що робить malware (MacSync)?": (
        "Викрадає паролі браузера, ключі доступу, криптогаманці, SSH-ключі, "
        "дані Keychain та інші конфіденційні файли."
    ),
    "Яким чином налагоджено захист?": (
        "Не виконувати невідомі Terminal-команди, використовувати лише офіційні джерела, "
        "EDR/антивірус, моніторинг shell-активності, блокування підозрілих сайтів і реклами."
    ),
    "Додаткові ризики": (
        "Атака не використовує експлойти — повністю залежить від дій користувача, "
        "що ускладнює виявлення та обходить класичні засоби захисту."
    ),
    "Посилання": "https://www.scworld.com/news/clickfix-campaigns-target-macos-users-via-macsync-infostealer"
}

# Таблиця
table = doc.add_table(rows=len(incident), cols=2)
table.style = 'Table Grid'

for i, (key, value) in enumerate(incident.items()):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = value

# Збереження
doc.save("Звіт_ClickFix_MacSync_macOS.docx")

print("DOCX файл успішно створено!")